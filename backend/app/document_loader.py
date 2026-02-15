from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader, UnstructuredFileLoader


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".xlsx", ".json"}


def _is_safe_zip_member(extract_root: Path, member: str) -> bool:
    destination = (extract_root / member).resolve()
    return str(destination).startswith(str(extract_root.resolve()))


def extract_if_zip(file_path: Path, extract_dir: Path) -> list[Path]:
    if file_path.suffix.lower() != ".zip":
        return [file_path]

    target_dir = extract_dir / file_path.stem
    if target_dir.exists():
        shutil.rmtree(target_dir)
    target_dir.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(file_path, "r") as zf:
        safe_members = [m for m in zf.namelist() if _is_safe_zip_member(target_dir, m)]
        for member in safe_members:
            zf.extract(member, target_dir)

    return [p for p in target_dir.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS]


def _load_text_like(path: Path) -> list[Document]:
    loader = TextLoader(str(path), autodetect_encoding=True)
    docs = loader.load()
    for d in docs:
        d.metadata.update({"source": str(path), "file_name": path.name, "extension": path.suffix.lower()})
    return docs


def _load_pdf(path: Path) -> list[Document]:
    docs = PyPDFLoader(str(path)).load()
    for d in docs:
        d.metadata.update({"source": str(path), "file_name": path.name, "extension": path.suffix.lower()})
    return docs


def _load_docx(path: Path) -> list[Document]:
    doc = DocxDocument(str(path))
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [
        Document(
            page_content=text,
            metadata={"source": str(path), "file_name": path.name, "extension": path.suffix.lower()},
        )
    ]


def _load_table(path: Path) -> list[Document]:
    if path.suffix.lower() == ".csv":
        df = pd.read_csv(path)
    else:
        df = pd.read_excel(path)
    return [
        Document(
            page_content=df.to_csv(index=False),
            metadata={"source": str(path), "file_name": path.name, "extension": path.suffix.lower()},
        )
    ]


def load_documents(path: Path) -> list[Document]:
    ext = path.suffix.lower()
    try:
        if ext in {".txt", ".md", ".json"}:
            return _load_text_like(path)
        if ext == ".pdf":
            return _load_pdf(path)
        if ext == ".docx":
            return _load_docx(path)
        if ext in {".csv", ".xlsx"}:
            return _load_table(path)

        docs = UnstructuredFileLoader(str(path)).load()
        for d in docs:
            d.metadata.update({"source": str(path), "file_name": path.name, "extension": ext})
        return docs
    except Exception:
        return _load_text_like(path)
